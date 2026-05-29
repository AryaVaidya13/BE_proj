import os
import json
import re
from typing import List, Dict

from groq import Groq
from core.config import GROQ_API_KEY

from reportlab.platypus import (
    BaseDocTemplate,
    Frame,
    PageTemplate,
    Paragraph,
    Spacer
)
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.pagesizes import A4
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont


try:
    pdfmetrics.registerFont(TTFont("TimesNewRoman", "times.ttf"))
    DEFAULT_FONT = "TimesNewRoman"
except:
    DEFAULT_FONT = "Helvetica"


class PaperAgentGroq:
    def __init__(
        self,
        literature: List[Dict],
        topic: str,
        experiments_bundle: Dict = None,
        output_dir: str = "outputs",
    ):
        self.literature = literature or []
        self.topic = topic or "Generated Paper"
        self.experiments_bundle = experiments_bundle or {}

        self.output_dir = os.path.abspath(output_dir)
        os.makedirs(self.output_dir, exist_ok=True)

        self.client = Groq(api_key=GROQ_API_KEY)

    def _safe_filename(self, text: str):
        return re.sub(r"[^A-Za-z0-9]+", "_", text).strip("_")

    def _generate_full_paper_with_llm(self):

        prompt = f"""
Write a formal SURVEY research paper.

Topic:
"{self.topic}"

Structure EXACTLY like this:

Abstract

Index Terms

I. Introduction

II. Related Work

III. Methodology

IV. Experimental Analysis

V. Discussion

VI. Conclusion and Future Work

Rules:
- Formal academic writing
- Paragraphs only
- No bullet points
- Cite papers like [1], [2]
- Do NOT hallucinate datasets or results

Literature:
{json.dumps(self.literature, indent=2)}

Experiments:
{json.dumps(self.experiments_bundle, indent=2)}

Return only the paper text.
"""

        response = self.client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.3
        )

        return response.choices[0].message.content.strip()

    def _generate_references(self):

        if not self.literature:
            return ["No references available"]

        refs = []

        for i, paper in enumerate(self.literature, start=1):
            title = paper.get("title", "Untitled")
            authors = ", ".join(paper.get("authors", [])) or "Unknown"
            year = paper.get("year", "n.d.")
            url = paper.get("url", "")

            ref = f"[{i}] {authors}, \"{title},\" {year}. Available: {url}"
            refs.append(ref)

        return refs

    def _add_page_number(self, canvas, doc):
        canvas.drawRightString(570, 20, str(doc.page))

    def _save_pdf(self, paper_text):

        safe_title = self._safe_filename(self.topic)
        pdf_path = os.path.join(self.output_dir, f"{safe_title}_IEEE_paper.pdf")

        styles = getSampleStyleSheet()

        base = ParagraphStyle(
            "Base",
            parent=styles["Normal"],
            fontName=DEFAULT_FONT,
            fontSize=10,
            leading=12,
            alignment=TA_JUSTIFY,
        )

        title_style = ParagraphStyle(
            "Title",
            parent=base,
            fontSize=18,
            alignment=TA_CENTER,
            spaceAfter=10,
        )

        author_style = ParagraphStyle(
            "Author",
            parent=base,
            alignment=TA_CENTER,
            spaceAfter=6,
        )

        heading_style = ParagraphStyle(
            "Heading",
            parent=base,
            fontSize=12,
            spaceBefore=10,
            spaceAfter=4,
        )

        doc = BaseDocTemplate(
            pdf_path,
            pagesize=A4,
            leftMargin=40,
            rightMargin=40,
            topMargin=50,
            bottomMargin=40,
        )

        frame1 = Frame(doc.leftMargin, doc.bottomMargin, doc.width / 2 - 6, doc.height)
        frame2 = Frame(doc.leftMargin + doc.width / 2 + 6, doc.bottomMargin, doc.width / 2 - 6, doc.height)

        template = PageTemplate(id="TwoCol", frames=[frame1, frame2], onPage=self._add_page_number)
        doc.addPageTemplates([template])

        story = []

        story.append(Paragraph(self.topic, title_style))

        story.append(Paragraph("Author Name 1", author_style))
        story.append(Paragraph("Department / University", author_style))
        story.append(Paragraph("email@domain.com", author_style))
        story.append(Spacer(1, 6))

        story.append(Paragraph("Author Name 2", author_style))
        story.append(Paragraph("Department / University", author_style))
        story.append(Paragraph("email@domain.com", author_style))
        story.append(Spacer(1, 14))

        lines = paper_text.split("\n")

        for line in lines:
            line = line.strip()

            if not line:
                story.append(Spacer(1, 6))
                continue

            if re.match(r"^(Abstract|Index Terms|I\.|II\.|III\.|IV\.|V\.|VI\.)", line):
                story.append(Paragraph(f"<b>{line}</b>", heading_style))
            else:
                story.append(Paragraph(line, base))

        story.append(Spacer(1, 10))
        story.append(Paragraph("<b>References</b>", heading_style))

        for ref in self._generate_references():
            story.append(Paragraph(ref, base))

        doc.build(story)

        print(f"📄 IEEE Paper saved → {pdf_path}")

        return pdf_path

    def _save_docx(self, paper_text):
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        safe_title = self._safe_filename(self.topic)
        docx_path = os.path.join(self.output_dir, f"{safe_title}_IEEE_paper.docx")

        doc = Document()

        # Title
        title = doc.add_heading(self.topic, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Authors
        authors_para = doc.add_paragraph()
        authors_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        authors_para.add_run("Author Name 1\nDepartment / University\nemail@domain.com\n\nAuthor Name 2\nDepartment / University\nemail@domain.com").bold = True

        # Content
        for line in paper_text.split("\n"):
            line = line.strip()
            if not line:
                continue

            if re.match(r"^(Abstract|Index Terms|I\.|II\.|III\.|IV\.|V\.|VI\.)", line):
                p = doc.add_heading(line, level=1)
                continue

            doc.add_paragraph(line)

        # References
        doc.add_heading("References", level=1)
        for ref in self._generate_references():
            doc.add_paragraph(ref)

        doc.save(docx_path)
        print(f"📄 Word Doc saved → {docx_path}")
        return docx_path

    def generate_paper(self):
        print("[📝] Generating paper...")

        paper_text = self._generate_full_paper_with_llm()

        txt_path = os.path.join(
            self.output_dir,
            f"{self._safe_filename(self.topic)}_paper.txt"
        )

        with open(txt_path, "w", encoding="utf-8") as f:
            f.write(paper_text)

        pdf_path = self._save_pdf(paper_text)
        docx_path = self._save_docx(paper_text)

        return {
            "text_path": txt_path,
            "pdf_path": pdf_path,
            "docx_path": docx_path
        }

    def run(self):
        return self.generate_paper()